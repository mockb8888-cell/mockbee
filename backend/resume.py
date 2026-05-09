  #!/usr/bin/env python3
"""
AI-Powered ATS Resume Builder
– User identified by email (no random IDs)
– Shared MongoDB via db.py  (same DB as mocker.py)
– Fixed PDF and DOCX alignment / layout issues
"""

import os, sys, json, subprocess, datetime

if sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')
# ── Auto-install missing packages ─────────────────────────────────────────────
def _ensure(pkg, import_as=None):
    try:
        __import__(import_as or pkg)
    except ImportError:
        print(f"  Installing {pkg}…")
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q",
                                "--break-system-packages"])

print("Checking dependencies…")
for pkg, imp in [("reportlab", None), ("python-docx", "docx"), ("requests", None), ("google-genai", "google.genai"), ("python-dotenv", "dotenv")]:
    _ensure(pkg, imp)
print("Dependencies ready.\n")

# ── MongoDB shared layer ──────────────────────────────────────────────────────
try:
    import db as _db
    MONGO_OK = True
except Exception as _me:
    MONGO_OK = False
    print(f"\033[93m  ⚠  MongoDB unavailable ({_me}). Resumes won't be saved to DB.\033[0m")

# ── ReportLab ─────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, KeepTogether)
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_JUSTIFY, TA_CENTER

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import requests

# ══════════════════════════════════════════════════════════════════════════════
#  ANSI Colours
# ══════════════════════════════════════════════════════════════════════════════
class C:
    RESET   = "\033[0m";  BOLD    = "\033[1m";  DIM    = "\033[2m"
    CYAN    = "\033[96m"; GREEN   = "\033[92m"; YELLOW = "\033[93m"
    BLUE    = "\033[94m"; MAGENTA = "\033[95m"; RED    = "\033[91m"
    WHITE   = "\033[97m"

def banner():
    print(f"""
{C.CYAN}{C.BOLD}
╔══════════════════════════════════════════════════════════════╗
║         🚀  AI-Powered ATS Resume Builder  🚀               ║
║         Generates Professional, ATS-Optimised Resumes        ║
╚══════════════════════════════════════════════════════════════╝
{C.RESET}""")

def section_header(title, icon=""):
    w = 60
    print(f"\n{C.BLUE}{C.BOLD}{'─'*w}{C.RESET}")
    print(f"{C.BLUE}{C.BOLD}  {icon}  {title}{C.RESET}")
    print(f"{C.BLUE}{C.BOLD}{'─'*w}{C.RESET}\n")

def prompt(label, required=True, hint=""):
    hint_s = f" {C.DIM}({hint}){C.RESET}" if hint else ""
    req_s  = (f" {C.RED}*{C.RESET}" if required
              else f" {C.DIM}(optional){C.RESET}")
    while True:
        val = input(f"  {C.WHITE}{label}{req_s}{hint_s}: {C.RESET}").strip()
        if val: return val
        if not required: return ""
        print(f"  {C.YELLOW}⚠  Required. Please enter a value.{C.RESET}")

def prompt_multiline(label, max_words=200, required=True):
    req_s = f" {C.RED}*{C.RESET}" if required else f" {C.DIM}(optional){C.RESET}"
    print(f"\n  {C.WHITE}{label}{req_s}{C.RESET}")
    print(f"  {C.DIM}(Press Enter twice when done. Max ~{max_words} words.){C.RESET}")
    lines = []
    while True:
        line = input("  ")
        if line == "" and lines: break
        lines.append(line)
    text = " ".join(lines).strip()
    words = text.split()
    if len(words) > max_words:
        text = " ".join(words[:max_words])
        print(f"  {C.YELLOW}⚠  Trimmed to {max_words} words.{C.RESET}")
    return text

def choose(label, options, allow_multi=False):
    print(f"\n  {C.WHITE}{label}{C.RESET}")
    for i, opt in enumerate(options, 1):
        print(f"    {C.CYAN}{i}.{C.RESET} {opt}")
    if allow_multi:
        print(f"  {C.DIM}(Comma-separated, e.g. 1,3){C.RESET}")
    while True:
        raw = input(f"  {C.WHITE}Choice: {C.RESET}").strip()
        if allow_multi:
            try:
                idxs = [int(x.strip()) for x in raw.split(",")]
                if all(1 <= i <= len(options) for i in idxs):
                    return [options[i-1] for i in idxs]
            except ValueError: pass
        else:
            try:
                idx = int(raw)
                if 1 <= idx <= len(options): return options[idx-1]
            except ValueError: pass
        print(f"  {C.YELLOW}⚠  Invalid choice.{C.RESET}")

# ══════════════════════════════════════════════════════════════════════════════
#  Data Collection
# ══════════════════════════════════════════════════════════════════════════════

def collect_email() -> str:
    section_header("IDENTIFY YOURSELF", "🔑")
    while True:
        val = input(f"  {C.WHITE}Your email address{C.RED} *{C.RESET}: ").strip().lower()
        if "@" in val and "." in val:
            return val
        print(f"  {C.YELLOW}⚠  Please enter a valid email.{C.RESET}")

def collect_personal_info():
    section_header("PERSONAL INFORMATION", "👤")
    return {
        "full_name":  prompt("Full Name"),
        "job_title":  prompt("Job Title / Desired Role",
                             hint="e.g. Senior Software Engineer"),
        "email":      prompt("Email Address"),
        "phone":      prompt("Phone Number",      required=False),
        "location":   prompt("City, Country",     required=False,
                             hint="e.g. London, UK"),
        "linkedin":   prompt("LinkedIn URL",      required=False),
        "github":     prompt("GitHub / Portfolio URL", required=False),
    }

def collect_work_experience():
    section_header("WORK EXPERIENCE", "💼")
    print(f"  {C.DIM}Add each role one at a time.{C.RESET}\n")
    experiences = []
    while True:
        print(f"  {C.GREEN}➕  Work Experience #{len(experiences)+1}{C.RESET}")
        exp = {
            "company":    prompt("Company Name"),
            "role":       prompt("Your Job Title"),
            "start_date": prompt("Start Date", hint="Jan 2020"),
            "end_date":   prompt("End Date",   hint="Dec 2023 or Present"),
            "summary":    prompt_multiline(
                              "Responsibilities & achievements", max_words=200),
        }
        experiences.append(exp)
        if input(f"\n  {C.CYAN}Add another? (y/n): {C.RESET}").strip().lower() != "y":
            break
    return experiences

def collect_education():
    section_header("EDUCATION", "🎓")
    educations = []
    while True:
        print(f"  {C.GREEN}➕  Education #{len(educations)+1}{C.RESET}")
        edu = {
            "institution": prompt("Institution / University"),
            "degree":      prompt("Degree / Qualification",
                                  hint="BSc Computer Science"),
            "year":        prompt("Year of Graduation", hint="2019"),
            "details":     prompt("Additional Details", required=False,
                                  hint="GPA, honours, modules"),
        }
        educations.append(edu)
        if input(f"\n  {C.CYAN}Add another? (y/n): {C.RESET}").strip().lower() != "y":
            break
    return educations

def collect_skills():
    section_header("SKILLS", "🛠")
    print(f"  {C.DIM}Enter skills separated by commas.{C.RESET}")
    raw = prompt("Skills", hint="Python, SQL, Leadership…")
    return [s.strip() for s in raw.split(",") if s.strip()]

def collect_summary():
    section_header("PROFESSIONAL SUMMARY", "📝")
    print(f"  {C.DIM}2-4 sentence overview of your professional identity.{C.RESET}")
    return prompt_multiline("Professional summary", max_words=100)

def collect_optional():
    section_header("ADDITIONAL INFORMATION", "✨")
    print(f"  {C.DIM}All fields optional.{C.RESET}\n")
    data = {}
    raw = prompt("Certifications", required=False, hint="AWS, PMP…")
    if raw: data["certifications"] = [c.strip() for c in raw.split(",") if c.strip()]
    raw = prompt("Languages", required=False, hint="English (Native), Spanish (B2)")
    if raw: data["languages"] = [l.strip() for l in raw.split(",") if l.strip()]
    if input(f"\n  {C.WHITE}Add notable projects? (y/n): {C.RESET}").strip().lower() == "y":
        data["projects"] = []
        while True:
            proj = {
                "name": prompt("Project Name"),
                "desc": prompt("Brief description", hint="1-2 sentences"),
                "link": prompt("Link (GitHub/Live)", required=False),
            }
            data["projects"].append(proj)
            if input(f"\n  {C.CYAN}Add another project? (y/n): {C.RESET}").strip().lower() != "y":
                break
    raw = prompt("Awards / Volunteering / Interests", required=False)
    if raw: data["other"] = raw
    return data


# ══════════════════════════════════════════════════════════════════════════════
#  AI Enhancement (Google Gemini)
# ══════════════════════════════════════════════════════════════════════════════
from google import genai as _genai
from google.genai import types as _genai_types
from dotenv import load_dotenv

# Try to load from sibling .env if available
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
_genai_client = _genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None

def ai_enhance(resume_data: dict) -> dict:
    if not GEMINI_API_KEY:
        print(f"  \033[93m⚠  GEMINI_API_KEY not found in environment. Skipping AI enhancement.\033[0m")
        return {}

    prompt_text = f"""You are an expert resume writer specialising in ATS-optimised content.
Given the following resume data in JSON, rewrite ONLY these keys to be more
professional, impactful, and ATS-friendly. Return ONLY valid JSON with these exact keys:
- "summary": Rewrite as a compelling 3-sentence professional summary.
- "experience_summaries": List of enhanced bullet descriptions (one per experience, \\n separated bullets).
- "skills_grouped": Group skills into dict {{category: [skill,…]}} or null.

Resume Data:
{json.dumps(resume_data, indent=2)}

Return ONLY a JSON object. No explanations, no markdown fences."""
    try:
        resp = _genai_client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt_text,
            config=_genai_types.GenerateContentConfig(
                response_mime_type="application/json",
            ),
        )
        text = resp.text.strip()
        text = text.lstrip("```json").lstrip("```").rstrip("```").strip()
        return json.loads(text)
    except Exception as e:
        print(f"  \033[91m⚠  Gemini AI Error: {e}\033[0m")
    return {}


# ══════════════════════════════════════════════════════════════════════════════
#  Colour Palettes
# ══════════════════════════════════════════════════════════════════════════════
PALETTES = {
    "Professional": {
        "primary":   colors.HexColor("#1a3c5e"),
        "accent":    colors.HexColor("#2e6da4"),
        "line":      colors.HexColor("#2e6da4"),
        "text":      colors.black,
        "font_name": "Times-Roman",
        "font_bold": "Times-Bold",
    },
    "Modern": {
        "primary":   colors.HexColor("#7c3aed"),
        "accent":    colors.HexColor("#a855f7"),
        "line":      colors.HexColor("#7c3aed"),
        "text":      colors.HexColor("#1f1f2e"),
        "font_name": "Helvetica",
        "font_bold": "Helvetica-Bold",
    },
    "Simple": {
        "primary":   colors.HexColor("#222222"),
        "accent":    colors.HexColor("#555555"),
        "line":      colors.HexColor("#cccccc"),
        "text":      colors.HexColor("#222222"),
        "font_name": "Helvetica",
        "font_bold": "Helvetica-Bold",
    },
    "Neat": {
        "primary":   colors.HexColor("#0f4c81"),
        "accent":    colors.HexColor("#1976d2"),
        "line":      colors.HexColor("#1976d2"),
        "text":      colors.HexColor("#222222"),
        "font_name": "Helvetica",
        "font_bold": "Helvetica-Bold",
    },
    "Black & White": {
        "primary":   colors.black,
        "accent":    colors.HexColor("#333333"),
        "line":      colors.black,
        "text":      colors.black,
        "font_name": "Helvetica",
        "font_bold": "Helvetica-Bold",
    },
}

DOCX_THEMES = {
    "Professional": {"header": (26, 60, 94),   "accent": (46, 109, 164)},
    "Modern":       {"header": (124, 58, 237),  "accent": (168, 85, 247)},
    "Simple":       {"header": (34, 34, 34),    "accent": (85, 85, 85)},
    "Neat":         {"header": (15, 76, 129),   "accent": (25, 118, 210)},
    "Black & White":{"header": (0, 0, 0),       "accent": (51, 51, 51)},
}


# ══════════════════════════════════════════════════════════════════════════════
#  PDF Generation  (fixed alignment)
# ══════════════════════════════════════════════════════════════════════════════
PAGE_W, PAGE_H = A4          # 595.28 × 841.89 pt
L_MARGIN = R_MARGIN = 0.65 * inch
CONTENT_W = PAGE_W - L_MARGIN - R_MARGIN

def _hex(colour):
    try:
        return f"{int(colour.red*255):02x}{int(colour.green*255):02x}{int(colour.blue*255):02x}"
    except Exception:
        return "000000"

def _styles(p):
    fn, fb = p["font_name"], p["font_bold"]
    return {
        "name": ParagraphStyle("name",
            fontName=fb, fontSize=22, leading=26,
            textColor=p["primary"], spaceAfter=2, alignment=TA_LEFT),

        "job_title_hdr": ParagraphStyle("job_title_hdr",
            fontName=fn, fontSize=13, leading=16,
            textColor=p["accent"], spaceAfter=3, alignment=TA_LEFT),

        "contact": ParagraphStyle("contact",
            fontName=fn, fontSize=8.5, leading=12,
            textColor=p["accent"], spaceAfter=4, alignment=TA_LEFT),

        "section": ParagraphStyle("section",
            fontName=fb, fontSize=10.5, leading=14,
            textColor=p["primary"], spaceBefore=8, spaceAfter=3,
            alignment=TA_LEFT),

        "job_role": ParagraphStyle("job_role",
            fontName=fb, fontSize=10, leading=13,
            textColor=p["text"], spaceAfter=1, alignment=TA_LEFT),

        "job_meta": ParagraphStyle("job_meta",
            fontName=fn, fontSize=9, leading=12,
            textColor=p["accent"], spaceAfter=3, alignment=TA_LEFT),

        "body": ParagraphStyle("body",
            fontName=fn, fontSize=9.5, leading=14,
            textColor=p["text"], spaceAfter=2, alignment=TA_JUSTIFY),

        # ── FIXED bullet: proper hanging indent so text wraps under text, not under bullet ──
        "bullet": ParagraphStyle("bullet",
            fontName=fn, fontSize=9.5, leading=14,
            textColor=p["text"], spaceAfter=2,
            leftIndent=14, firstLineIndent=0,    # hanging done via bullet char
            alignment=TA_LEFT),

        "skill": ParagraphStyle("skill",
            fontName=fn, fontSize=9.5, leading=13,
            textColor=p["text"], spaceAfter=2, alignment=TA_LEFT),

        "small": ParagraphStyle("small",
            fontName=fn, fontSize=8.5, leading=12,
            textColor=p["accent"], spaceAfter=2, alignment=TA_LEFT),
    }

def _hr(p, thick=0.7):
    return HRFlowable(width="100%", thickness=thick,
                      color=p["line"], spaceAfter=4, spaceBefore=2)

def _bullet_para(text: str, S: dict) -> Paragraph:
    """Render a single bullet line with a real bullet character."""
    clean = text.strip().lstrip("•‐-–").strip()
    if not clean:
        return None
    # Use XML bullet entity so ReportLab handles indent properly
    return Paragraph(f"\u2022\u00a0{clean}", S["bullet"])

def build_pdf_story(data: dict, p: dict) -> list:
    S = _styles(p)
    story = []
    pi  = data["personal"]
    opt = data.get("optional", {})

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(Paragraph(pi["full_name"], S["name"]))
    story.append(Paragraph(pi["job_title"], S["job_title_hdr"]))

    contacts = []
    if pi.get("email"):    contacts.append(pi["email"])
    if pi.get("phone"):    contacts.append(pi["phone"])
    if pi.get("location"): contacts.append(pi["location"])
    if pi.get("linkedin"): contacts.append(pi["linkedin"])
    if pi.get("github"):   contacts.append(pi["github"])
    if contacts:
        story.append(Paragraph("   |   ".join(contacts), S["contact"]))

    story.append(_hr(p, 1.5))

    # ── Summary ───────────────────────────────────────────────────────────────
    if data.get("summary"):
        story.append(Paragraph("PROFESSIONAL SUMMARY", S["section"]))
        story.append(_hr(p))
        story.append(Paragraph(data["summary"], S["body"]))
        story.append(Spacer(1, 5))

    # ── Experience ────────────────────────────────────────────────────────────
    if data.get("experience"):
        story.append(Paragraph("WORK EXPERIENCE", S["section"]))
        story.append(_hr(p))
        for exp in data["experience"]:
            block = []
            # role @ company on one line, dates on next — no inline colour hacks
            block.append(Paragraph(
                f"<b>{exp['role']}</b>  —  {exp['company']}", S["job_role"]))
            block.append(Paragraph(
                f"{exp['start_date']} – {exp['end_date']}", S["job_meta"]))
            for line in exp.get("summary", "").split("\n"):
                bp = _bullet_para(line, S)
                if bp: block.append(bp)
            block.append(Spacer(1, 5))
            story.append(KeepTogether(block))

    # ── Education ─────────────────────────────────────────────────────────────
    if data.get("education"):
        story.append(Paragraph("EDUCATION", S["section"]))
        story.append(_hr(p))
        for edu in data["education"]:
            block = []
            block.append(Paragraph(f"<b>{edu['degree']}</b>", S["job_role"]))
            meta = edu["institution"]
            if edu.get("year"): meta += f"  |  {edu['year']}"
            block.append(Paragraph(meta, S["job_meta"]))
            if edu.get("details"):
                block.append(Paragraph(edu["details"], S["small"]))
            block.append(Spacer(1, 5))
            story.append(KeepTogether(block))

    # ── Skills ────────────────────────────────────────────────────────────────
    if data.get("skills"):
        story.append(Paragraph("SKILLS", S["section"]))
        story.append(_hr(p))
        if data.get("skills_grouped") and isinstance(data["skills_grouped"], dict):
            for cat, items in data["skills_grouped"].items():
                story.append(Paragraph(
                    f"<b>{cat}:</b>\u00a0\u00a0{',  '.join(items)}", S["skill"]))
        else:
            # Wrap skills in rows of 6
            skills = data["skills"]
            rows = [skills[i:i+6] for i in range(0, len(skills), 6)]
            for row in rows:
                story.append(Paragraph(
                    "  \u2022  ".join(row), S["skill"]))
        story.append(Spacer(1, 4))

    # ── Certifications ────────────────────────────────────────────────────────
    if opt.get("certifications"):
        story.append(Paragraph("CERTIFICATIONS", S["section"]))
        story.append(_hr(p))
        for cert in opt["certifications"]:
            bp = _bullet_para(cert, S)
            if bp: story.append(bp)
        story.append(Spacer(1, 4))

    # ── Projects ──────────────────────────────────────────────────────────────
    if opt.get("projects"):
        story.append(Paragraph("PROJECTS", S["section"]))
        story.append(_hr(p))
        for proj in opt["projects"]:
            block = []
            block.append(Paragraph(f"<b>{proj['name']}</b>", S["job_role"]))
            block.append(Paragraph(proj["desc"], S["body"]))
            if proj.get("link"):
                block.append(Paragraph(proj["link"], S["small"]))
            block.append(Spacer(1, 4))
            story.append(KeepTogether(block))

    # ── Languages ─────────────────────────────────────────────────────────────
    if opt.get("languages"):
        story.append(Paragraph("LANGUAGES", S["section"]))
        story.append(_hr(p))
        story.append(Paragraph("  |  ".join(opt["languages"]), S["skill"]))
        story.append(Spacer(1, 4))

    # ── Other ─────────────────────────────────────────────────────────────────
    if opt.get("other"):
        story.append(Paragraph("ADDITIONAL INFORMATION", S["section"]))
        story.append(_hr(p))
        story.append(Paragraph(opt["other"], S["body"]))

    return story


def generate_pdf(data: dict, output_path: str) -> str:
    tname   = data.get("template", "Professional")
    palette = PALETTES.get(tname, PALETTES["Professional"])
    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=L_MARGIN, rightMargin=R_MARGIN,
        topMargin=0.55*inch, bottomMargin=0.55*inch,
    )
    doc.build(build_pdf_story(data, palette))
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX Generation  (fixed alignment)
# ══════════════════════════════════════════════════════════════════════════════

def _rgb(run, rgb):
    run.font.color.rgb = RGBColor(*rgb)

def _add_hr(doc, rgb=(0, 0, 0)):
    """Add a paragraph whose bottom border acts as a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    r, g, b = rgb
    bottom.set(qn("w:color"), f"{r:02x}{g:02x}{b:02x}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def _section_heading(doc, title: str, hc: tuple, ac: tuple):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    _rgb(run, hc)
    _add_hr(doc, ac)

def _body_para(doc, text: str, italic=False, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p

def _bullet_docx(doc, text: str, size=9.5):
    """
    Proper DOCX bullet using Word's built-in List Bullet style.
    Falls back to manual bullet if the style is unavailable.
    """
    text = text.strip().lstrip("•‐-–").strip()
    if not text:
        return
    try:
        p = doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)

def generate_docx(data: dict, output_path: str) -> str:
    tname = data.get("template", "Professional")
    theme = DOCX_THEMES.get(tname, DOCX_THEMES["Professional"])
    hc, ac = theme["header"], theme["accent"]

    doc = DocxDocument()

    # ── Page margins ──────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    # ── Default paragraph spacing (remove ugly gaps) ──────────────────────────
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(3)

    pi  = data["personal"]
    opt = data.get("optional", {})

    # ── Header ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(pi["full_name"])
    r.bold = True; r.font.size = Pt(22); _rgb(r, hc)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(pi["job_title"])
    r2.font.size = Pt(13); _rgb(r2, ac)

    contacts = [x for x in [
        pi.get("email"), pi.get("phone"), pi.get("location"),
        pi.get("linkedin"), pi.get("github")] if x]
    if contacts:
        pc = doc.add_paragraph()
        pc.paragraph_format.space_after = Pt(4)
        rc = pc.add_run("  |  ".join(contacts))
        rc.font.size = Pt(9); _rgb(rc, ac)

    _add_hr(doc, hc)

    # ── Summary ───────────────────────────────────────────────────────────────
    if data.get("summary"):
        _section_heading(doc, "Professional Summary", hc, ac)
        _body_para(doc, data["summary"], align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ── Experience ────────────────────────────────────────────────────────────
    if data.get("experience"):
        _section_heading(doc, "Work Experience", hc, ac)
        for exp in data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(f"{exp['role']}  —  {exp['company']}")
            r.bold = True; r.font.size = Pt(10)
            _body_para(doc, f"{exp['start_date']} – {exp['end_date']}",
                       italic=True, size=9)
            for line in exp.get("summary", "").split("\n"):
                _bullet_docx(doc, line)

    # ── Education ─────────────────────────────────────────────────────────────
    if data.get("education"):
        _section_heading(doc, "Education", hc, ac)
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(edu["degree"])
            r.bold = True; r.font.size = Pt(10)
            meta = edu["institution"]
            if edu.get("year"): meta += f"  |  {edu['year']}"
            _body_para(doc, meta, italic=True, size=9)
            if edu.get("details"):
                _body_para(doc, edu["details"], size=9)

    # ── Skills ────────────────────────────────────────────────────────────────
    if data.get("skills"):
        _section_heading(doc, "Skills", hc, ac)
        if data.get("skills_grouped") and isinstance(data["skills_grouped"], dict):
            for cat, items in data["skills_grouped"].items():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                rc = p.add_run(f"{cat}: ")
                rc.bold = True; rc.font.size = Pt(9.5)
                rv = p.add_run(", ".join(items))
                rv.font.size = Pt(9.5)
        else:
            _body_para(doc, ", ".join(data["skills"]))

    # ── Certifications ────────────────────────────────────────────────────────
    if opt.get("certifications"):
        _section_heading(doc, "Certifications", hc, ac)
        for cert in opt["certifications"]:
            _bullet_docx(doc, cert)

    # ── Projects ──────────────────────────────────────────────────────────────
    if opt.get("projects"):
        _section_heading(doc, "Projects", hc, ac)
        for proj in opt["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(proj["name"])
            r.bold = True; r.font.size = Pt(10)
            _body_para(doc, proj["desc"],
                       align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            if proj.get("link"):
                _body_para(doc, proj["link"], italic=True, size=9)

    # ── Languages ─────────────────────────────────────────────────────────────
    if opt.get("languages"):
        _section_heading(doc, "Languages", hc, ac)
        _body_para(doc, "  |  ".join(opt["languages"]))

    # ── Other ─────────────────────────────────────────────────────────────────
    if opt.get("other"):
        _section_heading(doc, "Additional Information", hc, ac)
        _body_para(doc, opt["other"], align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.save(output_path)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
#  Main Flow
# ══════════════════════════════════════════════════════════════════════════════
def main():
    banner()
    print(f"{C.GREEN}Welcome! Let's build your ATS-Optimised Resume step by step.{C.RESET}")
    print(f"{C.DIM}Fields marked {C.RED}*{C.DIM} are required.{C.RESET}\n")

    # ── Identify user by email ─────────────────────────────────────────────
    email = collect_email()
    if MONGO_OK:
        _db.upsert_user(email)

    data = {}
    data["email"]      = email
    data["personal"]   = collect_personal_info()
    # Pre-fill email in personal section if blank
    if not data["personal"].get("email"):
        data["personal"]["email"] = email

    data["summary"]    = collect_summary()
    data["experience"] = collect_work_experience()
    data["education"]  = collect_education()
    data["skills"]     = collect_skills()
    data["optional"]   = collect_optional()

    # ── Template ──────────────────────────────────────────────────────────────
    templates = [
        "Professional  – Classic serif, two-tone header",
        "Modern        – Bold purple accent, sans-serif",
        "Simple        – Minimal, clean single-column",
        "Neat          – Structured with ruled section dividers",
        "Black & White – Monochrome, printer-friendly, ultra-ATS safe",
    ]
    print()
    chosen = choose("Choose a Resume Template", templates)
    data["template"] = chosen.split("–")[0].strip()

    # ── AI Enhancement ────────────────────────────────────────────────────────
    print(f"\n{C.MAGENTA}🤖  Attempting AI content enhancement…{C.RESET}")
    print(f"{C.DIM}   (Requires GEMINI_API_KEY locally. Silently skipped if unavailable.){C.RESET}")
    enhanced = ai_enhance(data)
    if enhanced:
        print(f"{C.GREEN}✅  AI enhancement applied!{C.RESET}")
        if "summary" in enhanced:
            data["summary"] = enhanced["summary"]
        if "experience_summaries" in enhanced:
            for i, exp in enumerate(data["experience"]):
                if i < len(enhanced["experience_summaries"]):
                    exp["summary"] = enhanced["experience_summaries"][i]
        if enhanced.get("skills_grouped"):
            data["skills_grouped"] = enhanced["skills_grouped"]
    else:
        print(f"{C.DIM}   AI unavailable — using your original content.{C.RESET}")

    # ── Generate files ────────────────────────────────────────────────────────
    output_dir = os.path.join(os.path.expanduser("~"), "resumes")
    os.makedirs(output_dir, exist_ok=True)

    name_part = data["personal"]["full_name"].replace(" ", "_").replace("/", "-")
    ts        = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    base      = f"{output_dir}/{name_part}_Resume_{ts}"

    print(f"\n{C.MAGENTA}{'─'*60}{C.RESET}")
    print(f"{C.MAGENTA}{C.BOLD}  📄  Generating resume files…{C.RESET}")
    print(f"{C.MAGENTA}{'─'*60}{C.RESET}\n")

    pdf_path = docx_path = None

    try:
        pdf_path = generate_pdf(data, base + ".pdf")
        print(f"  {C.GREEN}✅  PDF  →  {pdf_path}{C.RESET}")
    except Exception as e:
        print(f"  {C.RED}❌  PDF failed: {e}{C.RESET}")

    try:
        docx_path = generate_docx(data, base + ".docx")
        print(f"  {C.GREEN}✅  DOCX →  {docx_path}{C.RESET}")
    except Exception as e:
        print(f"  {C.RED}❌  DOCX failed: {e}{C.RESET}")

    # ── Save raw JSON ──────────────────────────────────────────────────────────
    json_path = base + "_data.json"
    with open(json_path, "w") as f:
        json.dump(data, f, indent=2)
    print(f"  {C.DIM}💾  Data →  {json_path}{C.RESET}")

    # ── Persist to MongoDB ─────────────────────────────────────────────────────
    if MONGO_OK:
        try:
            _db.save_resume(email, data, pdf_path, docx_path)
            _db.archive_resume_snapshot(email, data, pdf_path, docx_path)
            print(f"  {C.GREEN}✅  Saved to MongoDB (user: {email}){C.RESET}")
        except Exception as e:
            print(f"  {C.YELLOW}⚠  MongoDB save failed: {e}{C.RESET}")

    print(f"\n{C.GREEN}{C.BOLD}  🎉  Done! Resume ready.{C.RESET}")
    print(f"  Template : {C.CYAN}{data['template']}{C.RESET}")
    print(f"  User ID  : {C.CYAN}{email}{C.RESET}")
    print(f"  Files    : {C.CYAN}{output_dir}{C.RESET}")

    # ── Open folder ───────────────────────────────────────────────────────────
    print()
    if input(f"  {C.WHITE}Open output folder now? (y/n): {C.RESET}").strip().lower() == "y":
        import platform
        try:
            plat = platform.system()
            if plat == "Darwin":    subprocess.Popen(["open", output_dir])
            elif plat == "Windows": subprocess.Popen(["explorer", output_dir])
            else:                   subprocess.Popen(["xdg-open", output_dir])
        except Exception:
            print(f"  {C.DIM}Could not open folder automatically.{C.RESET}")

    print(f"\n{C.DIM}  ATS Tips:")
    print(f"  • Match keywords from the job description in your summary & skills.")
    print(f"  • Use clean filename: FirstName_LastName_Resume.pdf")
    print(f"  • Submit PDF for most roles; DOCX where explicitly required.{C.RESET}\n")


if __name__ == "__main__":
    main()
